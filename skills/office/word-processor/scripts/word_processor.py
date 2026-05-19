#!/usr/bin/env python3
"""
Word 文档生成器

用法:
    python word_processor.py --title "项目报告" --output report.docx
    python word_processor.py --title "合同" --output contract.docx --sections title,content,table
    python word_processor.py --dry-run --title "测试"
"""

import sys
import json
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def create_document(title="", sections=None, output_path=None, dry_run=False):
    """创建 Word 文档"""
    if not HAS_DOCX:
        return {"error": "python-docx 未安装，请运行: pip install python-docx"}

    if dry_run:
        return {"status": "dry_run", "title": title, "sections": sections or [], "message": "仅模拟运行，未生成文件"}

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    section_configs = sections or ["title", "content"]
    created_sections = []

    for i, section_type in enumerate(section_configs):
        if section_type == "title":
            add_title_page(doc, title or "文档")
            created_sections.append({"index": i, "type": "title", "title": title})
        elif section_type == "content":
            add_content_section(doc, f"第 {i} 节")
            created_sections.append({"index": i, "type": "content"})
        elif section_type == "table":
            add_table_section(doc, f"表格 {i}")
            created_sections.append({"index": i, "type": "table"})
        elif section_type == "list":
            add_list_section(doc, f"列表 {i}")
            created_sections.append({"index": i, "type": "list"})

    if output_path:
        doc.save(output_path)
        return {
            "status": "success",
            "output": str(Path(output_path).resolve()),
            "sections": len(created_sections),
            "section_details": created_sections,
        }
    else:
        return {"error": "请指定 --output 参数"}


def add_title_page(doc, title_text):
    """添加标题页"""
    # 添加空行
    for _ in range(6):
        doc.add_paragraph("")

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("自动生成文档")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 分页
    doc.add_page_break()


def add_content_section(doc, section_title):
    """添加内容节"""
    # 一级标题
    heading = doc.add_heading(section_title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

    # 段落
    doc.add_paragraph("在此添加正文内容。可以包含多个段落，每个段落之间用空行分隔。")

    # 二级标题
    sub_heading = doc.add_heading("子标题", level=2)
    for run in sub_heading.runs:
        run.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)

    doc.add_paragraph("详细内容...")


def add_table_section(doc, section_title):
    """添加表格节"""
    heading = doc.add_heading(section_title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

    # 创建表格
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ["项目", "数值", "备注"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # 设置表头背景色（通过 shading）
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "2C5F8A")
        shading_elm.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # 示例数据
    data = [
        ["项目 A", "100", "已完成"],
        ["项目 B", "200", "进行中"],
        ["项目 C", "150", "待开始"],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_text in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_text


def add_list_section(doc, section_title):
    """添加列表节"""
    heading = doc.add_heading(section_title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

    # 无序列表
    doc.add_paragraph("无序列表项 1", style="List Bullet")
    doc.add_paragraph("无序列表项 2", style="List Bullet")
    doc.add_paragraph("无序列表项 3", style="List Bullet")

    # 有序列表
    doc.add_paragraph("有序列表项 1", style="List Number")
    doc.add_paragraph("有序列表项 2", style="List Number")
    doc.add_paragraph("有序列表项 3", style="List Number")


def main():
    parser = argparse.ArgumentParser(description="Word 文档生成器")
    parser.add_argument("--title", default="文档", help="文档标题")
    parser.add_argument("--output", "-o", required=False, help="输出文件路径 (.docx)")
    parser.add_argument(
        "--sections",
        default="title,content,table",
        help="章节类型列表 (title,content,table,list)，逗号分隔",
    )
    parser.add_argument("--dry-run", action="store_true", help="仅模拟运行")
    parser.add_argument("--json", action="store_true", help="输出 JSON 格式")

    args = parser.parse_args()
    sections = [s.strip() for s in args.sections.split(",")]

    result = create_document(
        title=args.title,
        sections=sections,
        output_path=args.output,
        dry_run=args.dry_run,
    )

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        if result.get("error"):
            print(f"❌ {result['error']}", file=sys.stderr)
            sys.exit(1)
        elif result.get("status") == "dry_run":
            print(f"📄 模拟生成: {result['title']}")
            print(f"   章节: {', '.join(result['sections'])}")
        else:
            print(f"✅ 已生成: {result['output']}")
            print(f"   共 {result['sections']} 个章节")


if __name__ == "__main__":
    main()
